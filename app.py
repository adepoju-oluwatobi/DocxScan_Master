import os
from docx import Document
from flask import Flask, render_template, request


UPLOAD_FOLDER = 'uploads'  # Define the upload folder
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


@app.route('/')
def homepage():
    return render_template('upload.html')


def find_duplicate_questions(doc):
    question_dict = {}
    duplicates = {}

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text and text[0].isdigit():  # Check if the line starts with a number
            question_number, question_text = text.split('.', 1)
            question_number = question_number.strip()
            question_text = question_text.strip()

            if question_text in question_dict:
                duplicates.setdefault(question_text, []).append(question_number)
            else:
                question_dict[question_text] = question_number

    return question_dict, duplicates


def get_font_styles(doc):
    font_styles = set()

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font_styles.add(run.font.name)

    return list(font_styles)


def get_font_sizes(doc):
    font_sizes = set()

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font_sizes.add(run.font.size)

    return list(font_sizes)


def get_line_spacing(doc):
    line_spacing = set()

    for paragraph in doc.paragraphs:
        line_spacing.add(paragraph.paragraph_format.line_spacing)

    return list(line_spacing)


@app.route('/upload', methods=['POST'])
def upload_file():
    if request.method == 'POST':
        uploaded_file = request.files['file']
        if uploaded_file.filename != '':
            if uploaded_file.filename.endswith('.docx'):
                # Save the uploaded file to the server
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], uploaded_file.filename)
                uploaded_file.save(file_path)

                # Process the uploaded DOCX file
                doc = Document(file_path)

                # Get duplicate questions
                question_dict, duplicates = find_duplicate_questions(doc)

                # Extract duplicate questions with numbers
                duplicate_questions = []
                for question_text, numbers in duplicates.items():
                    initial_number = question_dict.get(question_text)  # Get the initial number
                    duplicate_numbers = ", ".join(numbers)  # Join duplicate question numbers
                    duplicate_questions.append(f"{initial_number}. {question_text} - Duplicates: {duplicate_numbers}")

                # Check if there are no duplicate questions
                if not duplicate_questions:
                    duplicate_questions.append("No duplicate questions")

                # Get font styles, sizes, and line spacing
                font_styles = get_font_styles(doc)
                font_sizes = get_font_sizes(doc)
                line_spacing = get_line_spacing(doc)

                # Return the results to the frontend
                return render_template('results.html', duplicate_questions=duplicate_questions,
                                       font_styles=font_styles,
                                       font_sizes=font_sizes,
                                       line_spacing=line_spacing)
            else:
                err_msg = "Please upload a .docx file."
                return render_template('upload.html', err_msg=err_msg)
        else:
            err_msg = "No file selected."
            return render_template('upload.html', err_msg=err_msg)


if __name__ == "__main__":
    # Start Flask in a separate thread
    app.run(debug=True)
